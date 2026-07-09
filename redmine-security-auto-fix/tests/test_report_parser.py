import sys
import tempfile
import unittest
from pathlib import Path

SCRIPTS = Path(__file__).resolve().parents[1] / "scripts"
sys.path.insert(0, str(SCRIPTS))

from report_parser import parse_report


class ReportParserTests(unittest.TestCase):
    def test_csv_aliases_are_normalized(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "report.csv"
            path.write_text(
                "漏洞名称,风险等级,漏洞描述,整改建议,漏洞地址\n"
                "CORS配置不当,中危,来源校验不足,配置域名白名单,https://example.test/a\n",
                encoding="utf-8",
            )
            result = parse_report(path)
        self.assertEqual(result["total"], 1)
        vuln = result["vulns"][0]
        self.assertEqual(vuln["name"], "CORS配置不当")
        self.assertEqual(vuln["level"], "medium")
        self.assertEqual(vuln["fix_suggestion"], "配置域名白名单")
        self.assertEqual(vuln["urls"], ["https://example.test/a"])

    def test_json_items_are_normalized(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "report.json"
            path.write_text(
                '{"items":[{"title":"弱口令","severity":"high","remediation":"修改密码策略"}]}',
                encoding="utf-8",
            )
            result = parse_report(path)
        self.assertEqual(result["vulns"][0]["level"], "high")
        self.assertEqual(result["vulns"][0]["fix_suggestion"], "修改密码策略")

    def test_docx_heading_styles_do_not_split_body_text(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "report.docx"
            document = Document()
            document.add_heading("【高危】SQL注入", level=2)
            document.add_paragraph("漏洞描述：")
            document.add_paragraph("攻击者利用注入漏洞读取数据库信息。")
            document.add_paragraph("测试过程：")
            document.add_paragraph("该接口存在SQL注入漏洞。")
            document.add_paragraph("加固建议：")
            document.add_paragraph("使用参数化查询。")
            document.add_heading("【中危】详细的报错信息*2", level=2)
            document.add_paragraph("漏洞危害：")
            document.add_paragraph("错误信息可能泄露数据库连接信息。")
            document.save(path)

            result = parse_report(path)

        self.assertEqual(result["total"], 2)
        self.assertEqual(result["vulns"][0]["name"], "SQL注入")
        self.assertIn("该接口存在SQL注入漏洞", result["vulns"][0]["description"])
        self.assertEqual(result["vulns"][0]["fix_suggestion"], "使用参数化查询。")
        self.assertEqual(result["vulns"][1]["name"], "详细的报错信息*2")


if __name__ == "__main__":
    unittest.main()
