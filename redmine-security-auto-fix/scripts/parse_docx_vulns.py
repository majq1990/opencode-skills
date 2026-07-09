#!/usr/bin/env python
# -*- coding: utf-8 -*-

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def extract_full_vuln_details(doc):
    """从段落中完整提取漏洞详情（描述、危害、测试过程、加固建议）"""
    vulns = []
    vuln_pattern = re.compile(r'【(严重|高危|中危|低危|信息)】(.+?)(?:\*(\d+))?$')
    
    level_map = {'严重': 'critical', '高危': 'high', '中危': 'medium', '低危': 'low', '信息': 'info'}
    
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    
    current_vuln = None
    current_section = None
    section_content = []
    
    for i, text in enumerate(paragraphs):
        if not text:
            continue
        
        match = vuln_pattern.match(text)
        if match:
            if current_vuln:
                if current_section and section_content:
                    current_vuln[current_section] = '\n'.join(section_content).strip()
                vulns.append(current_vuln)
            
            level_cn = match.group(1)
            name = match.group(2).strip()
            count = int(match.group(3)) if match.group(3) else 1
            
            current_vuln = {
                "name": name,
                "level": level_map.get(level_cn, 'medium'),
                "count": count,
                "description": "",
                "harm": "",
                "test_process": "",
                "fix_suggestion": "",
                "urls": []
            }
            current_section = None
            section_content = []
            continue
        
        if current_vuln is None:
            continue
        
        if text == '漏洞描述：':
            if current_section and section_content:
                current_vuln[current_section] = '\n'.join(section_content).strip()
            current_section = 'description'
            section_content = []
        elif text == '漏洞危害：':
            if current_section and section_content:
                current_vuln[current_section] = '\n'.join(section_content).strip()
            current_section = 'harm'
            section_content = []
        elif text == '漏洞简述：':
            if current_section and section_content:
                current_vuln[current_section] = '\n'.join(section_content).strip()
            current_section = None
            section_content = []
        elif text == '测试过程：':
            if current_section and section_content:
                current_vuln[current_section] = '\n'.join(section_content).strip()
            current_section = 'test_process'
            section_content = []
        elif text == '加固建议：':
            if current_section and section_content:
                current_vuln[current_section] = '\n'.join(section_content).strip()
            current_section = 'fix_suggestion'
            section_content = []
        elif current_section:
            section_content.append(text)
    
    if current_vuln:
        if current_section and section_content:
            current_vuln[current_section] = '\n'.join(section_content).strip()
        vulns.append(current_vuln)
    
    return vulns


def extract_urls_from_tables(doc):
    """从表格中提取URL列表"""
    urls = []
    
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.columns) < 2:
            continue
        
        first_cell = table.rows[0].cells[0].text.strip()
        if '漏洞信息' not in first_cell and '漏洞' not in first_cell:
            continue
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                key = cells[0]
                value = cells[1]
                if 'URL' in key or '地址' in key:
                    url_match = re.search(r'(https?://[^\s]+)', value)
                    if url_match:
                        urls.append(url_match.group(1))
    
    return urls


def parse_vuln_report(docx_path: str, issue_id: int) -> dict:
    doc = Document(docx_path)
    
    vulns = extract_full_vuln_details(doc)
    table_urls = extract_urls_from_tables(doc)
    
    url_index = 0
    for vuln in vulns:
        count = vuln.get('count', 1)
        vuln_urls = []
        for _ in range(count):
            if url_index < len(table_urls):
                vuln_urls.append(table_urls[url_index])
                url_index += 1
        vuln['urls'] = vuln_urls
    
    expanded_vulns = []
    vuln_id = 1
    
    for vuln in vulns:
        count = vuln.get('count', 1)
        for i in range(count):
            entry = {
                "id": vuln_id,
                "name": vuln['name'],
                "level": vuln['level'],
                "description": vuln.get('description', ''),
                "harm": vuln.get('harm', ''),
                "test_process": vuln.get('test_process', ''),
                "fix_suggestion": vuln.get('fix_suggestion', ''),
                "urls": vuln['urls'][i:i+1] if i < len(vuln['urls']) else []
            }
            expanded_vulns.append(entry)
            vuln_id += 1
    
    result = {
        "issue_id": issue_id,
        "source": "paragraphs",
        "vulns": expanded_vulns,
        "total": len(expanded_vulns),
        "stats": {
            "critical": sum(1 for v in expanded_vulns if v['level'] == 'critical'),
            "high": sum(1 for v in expanded_vulns if v['level'] == 'high'),
            "medium": sum(1 for v in expanded_vulns if v['level'] == 'medium'),
            "low": sum(1 for v in expanded_vulns if v['level'] == 'low'),
            "info": sum(1 for v in expanded_vulns if v['level'] == 'info')
        }
    }
    
    return result


def main():
    parser = argparse.ArgumentParser(description='Parse vulnerability DOCX report')
    parser.add_argument('docx_path', help='Path to the DOCX vulnerability report')
    parser.add_argument('--issue-id', type=int, required=True, help='Redmine issue ID')
    parser.add_argument('--output', help='Output JSON file path')
    
    args = parser.parse_args()
    
    if not Path(args.docx_path).exists():
        print(f"ERROR: File not found: {args.docx_path}")
        sys.exit(1)
    
    result = parse_vuln_report(args.docx_path, args.issue_id)
    
    output_path = args.output or f"{Path(args.docx_path).parent}/{args.issue_id}_vulns.json"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    
    print(f"Parsed {result['total']} vulnerabilities (source: {result['source']})")
    print(f"Stats: critical={result['stats']['critical']}, high={result['stats']['high']}, medium={result['stats']['medium']}, low={result['stats']['low']}")
    print(f"Output saved to: {output_path}")


if __name__ == '__main__':
    main()
