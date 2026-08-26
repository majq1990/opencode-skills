#!/usr/bin/env python3
"""Normalize heterogeneous vulnerability reports into one JSON schema."""

from __future__ import annotations

import csv
import json
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any

LEVELS = {
    "严重": "critical",
    "致命": "critical",
    "超危": "critical",
    "critical": "critical",
    "高危": "high",
    "高风险": "high",
    "high": "high",
    "中危": "medium",
    "中风险": "medium",
    "medium": "medium",
    "低危": "low",
    "低风险": "low",
    "low": "low",
    "信息": "info",
    "提示": "info",
    "info": "info",
}

ALIASES = {
    "name": (
        "漏洞名称",
        "漏洞名",
        "风险名称",
        "问题名称",
        "组件名称",
        "缺陷类型",
        "漏洞类型",
        "标题",
        "name",
        "title",
        "category",
    ),
    "level": (
        "风险等级",
        "风险级别",
        "组件等级",
        "漏洞等级",
        "危险等级",
        "级别",
        "severity",
        "risk",
        "level",
    ),
    "description": (
        "漏洞描述",
        "问题描述",
        "风险描述",
        "漏洞简述",
        "描述",
        "组件来源",
        "问题所在文件",
        "description",
        "detail",
    ),
    "harm": ("漏洞危害", "风险影响", "影响", "危害", "impact", "harm"),
    "fix_suggestion": (
        "加固建议",
        "修复建议",
        "整改建议",
        "解决方案",
        "修复方案",
        "recommendation",
        "remediation",
        "solution",
    ),
    "url": ("漏洞地址", "涉及url", "url", "uri", "位置", "路径"),
    "cve": ("cve", "cve编号"),
    "cwe": ("cwe", "cwe编号"),
}


def _clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _find_key(row: dict, aliases: tuple[str, ...]) -> str:
    lowered = {_clean(k).lower(): v for k, v in row.items()}
    for alias in aliases:
        if alias.lower() in lowered:
            return _clean(lowered[alias.lower()])
    return ""


def _normalize_level(value: str) -> str:
    text = _clean(value).lower()
    for key, level in LEVELS.items():
        if key.lower() in text:
            return level
    # Single-character Chinese severity (e.g. "高", "中", "低")
    single = {
        "严重": "critical",
        "高": "high",
        "中": "medium",
        "低": "low",
        "信息": "info",
    }
    for char, level in single.items():
        if char in text:
            return level
    return "medium"


def normalize_rows(rows: list[dict], source_file: str) -> list[dict]:
    vulns = []
    for row in rows:
        name = _find_key(row, ALIASES["name"])
        if not name:
            name = _find_key(row, ("描述",))
        if not name:
            rule = _find_key(row, ("扫描规则",))
            finding = _find_key(row, ("扫描结果",))
            if rule or finding:
                name = f"代码扫描发现：{finding or rule[:80]}"
        if not name:
            continue
        urls = _find_key(row, ALIASES["url"])
        vulns.append(
            {
                "name": name,
                "level": _normalize_level(_find_key(row, ALIASES["level"])),
                "description": _find_key(row, ALIASES["description"]),
                "harm": _find_key(row, ALIASES["harm"]),
                "fix_suggestion": _find_key(row, ALIASES["fix_suggestion"]),
                "urls": re.findall(r"https?://[^\s,;]+", urls),
                "cve": _find_key(row, ALIASES["cve"]),
                "cwe": _find_key(row, ALIASES["cwe"]),
                "source_file": source_file,
            }
        )
    return vulns


def _parse_delimited(path: Path) -> list[dict]:
    for encoding in ("utf-8-sig", "gb18030", "utf-8"):
        try:
            with path.open("r", encoding=encoding, newline="") as handle:
                sample = handle.read(4096)
                handle.seek(0)
                dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
                return list(csv.DictReader(handle, dialect=dialect))
        except (UnicodeDecodeError, csv.Error):
            continue
    raise ValueError(f"Unable to decode delimited report: {path}")


def _parse_excel(path: Path) -> list[dict]:
    try:
        import pandas as pd
    except ImportError as exc:
        raise RuntimeError("Excel parsing requires pandas and openpyxl/xlrd") from exc
    sheets = pd.read_excel(path, sheet_name=None, header=None, dtype=str)
    rows = []
    for sheet_name, frame in sheets.items():
        frame = frame.fillna("")
        if frame.empty or len(frame.columns) == 0:
            continue
        header_index = 0
        best_score = -1
        aliases = {
            alias.lower()
            for values in ALIASES.values()
            for alias in values
        } | {"扫描规则", "扫描结果", "整改建议", "组件名称", "风险名称"}
        for index in range(min(20, len(frame))):
            values = [_clean(value).lower() for value in frame.iloc[index].tolist()]
            score = sum(1 for value in values if value in aliases)
            if score > best_score:
                header_index = index
                best_score = score
        headers = [
            _clean(value) or f"column_{column}"
            for column, value in enumerate(frame.iloc[header_index].tolist())
        ]
        data = frame.iloc[header_index + 1 :].copy()
        data.columns = headers
        for row in data.to_dict(orient="records"):
            row["_sheet"] = sheet_name
            rows.append(row)
    return rows


def _merge_heading_with_tables(heading_rows: list[dict], table_rows: list[dict]) -> list[dict]:
    if not heading_rows or not table_rows:
        return heading_rows or table_rows

    def bigrams(text: str) -> set[str]:
        chars = re.sub(r"[\d.\-a-z_/\\]+", "", text or "", flags=re.I)
        return {chars[i:i+2] for i in range(len(chars) - 1)}

    def keywords(text: str) -> set[str]:
        return set(re.findall(r"[\u4e00-\u9fff]{2,}", text or ""))

    # Build a score matrix between every heading and every table, then assign
    # greedily by descending score so a strong match is never stolen by an
    # earlier heading that only weakly overlaps the same table.
    scores: list[tuple[float, int, int]] = []
    heading_bigrams = [(bigrams(h.get("漏洞名称") or ""), keywords(h.get("漏洞名称") or "")) for h in heading_rows]
    table_bigrams = [(bigrams(t.get("漏洞名称") or ""), keywords(t.get("漏洞名称") or "")) for t in table_rows]
    for hi, (hb, hk) in enumerate(heading_bigrams):
        hname = heading_rows[hi].get("漏洞名称") or ""
        for ti, (tb, tk) in enumerate(table_bigrams):
            tname = table_rows[ti].get("漏洞名称") or ""
            score = len(hb & tb) * 1.5 + len(hk & tk) * 2.0
            if hname in tname or tname in hname:
                score += 10
            scores.append((score, hi, ti))

    heading_to_table: dict[int, int] = {}
    used_tables: set[int] = set()
    for score, hi, ti in sorted(scores, key=lambda x: x[0], reverse=True):
        if score < 2 or hi in heading_to_table or ti in used_tables:
            continue
        heading_to_table[hi] = ti
        used_tables.add(ti)

    result: list[dict] = []
    for hi, heading in enumerate(heading_rows):
        ti = heading_to_table.get(hi)
        if ti is None:
            result.append(heading)
            continue
        table_row = table_rows[ti]
        merged = dict(heading)
        for field in ("漏洞描述", "漏洞危害", "加固建议", "漏洞地址", "风险等级"):
            if not merged.get(field) and table_row.get(field):
                merged[field] = table_row[field]
        result.append(merged)
    return result


def _parse_docx(path: Path) -> list[dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX parsing requires python-docx") from exc

    doc = Document(path)
    rows = []
    for table in doc.tables:
        if not table.rows:
            continue
        vertical = {}
        vertical_keys = {
            "风险名称": "漏洞名称",
            "漏洞名称": "漏洞名称",
            "问题名称": "漏洞名称",
            "检测项": "漏洞名称",
            "检测目": "漏洞名称",
            "检测内容": "漏洞名称",
            "风险级别": "风险等级",
            "风险等级": "风险等级",
            "漏洞等级": "风险等级",
            "风险描述": "漏洞描述",
            "漏洞描述": "漏洞描述",
            "漏洞危害": "漏洞危害",
            "风险影响": "漏洞危害",
            "风险分析": "漏洞危害",
            "结果描述": "漏洞危害",
            "加固建议": "加固建议",
            "修复建议": "加固建议",
            "整改建议": "加固建议",
            "解决方案": "加固建议",
            "漏洞链接": "漏洞地址",
            "涉及URL": "漏洞地址",
        }
        for table_row in table.rows:
            values = [_clean(cell.text) for cell in table_row.cells]
            if len(values) >= 2:
                label = values[0]
                matched = next(
                    (v for k, v in vertical_keys.items() if k in label),
                    None,
                )
                if matched:
                    vertical[matched] = values[1]
        if vertical.get("漏洞名称"):
            rows.append(vertical)
            continue
        headers = [_clean(cell.text) for cell in table.rows[0].cells]
        if not any(headers):
            continue
        for table_row in table.rows[1:]:
            values = [_clean(cell.text) for cell in table_row.cells]
            rows.append(dict(zip(headers, values)))

    paragraph_records = [
        (paragraph.text.strip(), paragraph.style.name or "")
        for paragraph in doc.paragraphs
        if paragraph.text.strip()
    ]
    heading = re.compile(
        r"(?:【(?P<level>严重|高危|中危|低危|信息)】|\[(?P<level_en>critical|high|medium|low|info)\])?\s*"
        r"(?P<name>[^\n：:]{3,100})(?:\*(?P<count>\d+))?$",
        re.I,
    )
    explicit_heading = re.compile(
        r"^(?:【(?P<level>严重|高危|中危|低危|信息)】|\[(?P<level_en>critical|high|medium|low|info)\])\s*"
        r"(?P<name>.+)$",
        re.I,
    )
    section_re = re.compile(
        r"^(漏洞描述|问题描述|漏洞危害|影响|漏洞简述|测试过程|加固建议|修复建议|整改建议|解决方案)[：:]?\s*(.*)$",
        re.I,
    )
    has_styled_headings = any(
        style.lower().startswith("heading") and explicit_heading.match(line)
        for line, style in paragraph_records
    )
    current = None
    current_section = None
    for line, style in paragraph_records:
        section = section_re.match(line)
        if section and current:
            label, content = section.groups()
            current_section = (
                "fix_suggestion"
                if any(x in label for x in ("建议", "方案"))
                else "harm"
                if any(x in label for x in ("危害", "影响"))
                else "description"
            )
            if content:
                current[current_section] = content
            continue
        match = explicit_heading.match(line) if has_styled_headings else heading.match(line)
        is_heading = bool(
            match
            and (
                has_styled_headings
                and style.lower().startswith("heading")
                or not has_styled_headings
                and (
                    match.group("level")
                    or match.group("level_en")
                    or any(
                        word in line
                        for word in ("漏洞", "注入", "跨站", "越权", "泄露")
                    )
                )
            )
        )
        if is_heading:
            if current:
                rows.append(current)
            current = {
                "漏洞名称": match.group("name"),
                "风险等级": match.group("level") or match.group("level_en") or "",
                "漏洞描述": "",
                "漏洞危害": "",
                "加固建议": "",
                "source": "paragraph",
            }
            current_section = None
        elif current and current_section:
            key = {
                "description": "漏洞描述",
                "harm": "漏洞危害",
                "fix_suggestion": "加固建议",
            }[current_section]
            current[key] = _clean(f"{current.get(key, '')} {line}")
    if current:
        rows.append(current)

    if rows:
        table_rows = [r for r in rows if r.get("漏洞名称") and (r.get("source") != "paragraph")]
        para_rows = [r for r in rows if r.get("source") == "paragraph"]
        rows = _merge_heading_with_tables(para_rows, table_rows)
    return rows


def _parse_pdf(path: Path) -> list[dict]:
    rows = []
    text_parts = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        text_parts = [(page.extract_text() or "") for page in reader.pages]
    except ImportError:
        pass
    except Exception:
        text_parts = []

    if not any(part.strip() for part in text_parts):
        try:
            import pdfplumber
        except ImportError as exc:
            raise RuntimeError("PDF parsing requires pypdf or pdfplumber") from exc
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
                for table in page.extract_tables() or []:
                    if len(table) < 2:
                        continue
                    headers = [_clean(cell) for cell in table[0]]
                    rows.extend(dict(zip(headers, row)) for row in table[1:])
    text = "\n".join(text_parts)
    specialized = _parse_fortify_pdf(text)
    if specialized:
        return specialized
    cve_bulletin = _parse_cve_bulletin(text)
    if cve_bulletin:
        return cve_bulletin
    if rows:
        normalized = normalize_rows(rows, path.name)
        if normalized:
            return rows
    return _parse_labeled_text(text)


def _parse_fortify_pdf(text: str) -> list[dict]:
    matches = list(
        re.finditer(
            r"Category:\s*(?P<name>.+?)\s*\((?P<count>\d+)\s+Issues?\)",
            text,
            re.I,
        )
    )
    rows = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        block = text[match.end() : end]
        level_match = re.search(
            r"Fortify Priority:\s*(Critical|High|Medium|Low)", block, re.I
        )
        recommendation = ""
        recommendation_match = re.search(
            r"(?:Recommendations?|Recommendation Summary)\s*[:\n]\s*(.*?)(?=\n(?:[A-Z][A-Za-z ]{2,}:|Category:)|\Z)",
            block,
            re.S,
        )
        if recommendation_match:
            recommendation = _clean(recommendation_match.group(1))[:3000]
        rows.append(
            {
                "漏洞名称": _clean(match.group("name")),
                "风险等级": level_match.group(1) if level_match else "",
                "漏洞描述": _clean(block[:1500]),
                "加固建议": recommendation,
            }
        )
    return rows


def _parse_cve_bulletin(text: str) -> list[dict]:
    cves = sorted(set(re.findall(r"CVE-\d{4}-\d{4,7}", text, re.I)))
    if not cves:
        return []
    title_lines = [
        _clean(line)
        for line in text.splitlines()[:20]
        if _clean(line) and not re.fullmatch(r"[\d年月日、,，()（）\s-]+", _clean(line))
    ]
    name = " ".join(title_lines[:4])
    if len(name) > 180:
        name = f"{'、'.join(cves)} 安全漏洞"
    suggestion = ""
    match = re.search(
        r"(?:处置建议|修复建议|解决方案)\s*(.*?)(?=\n第[五六七八九十]章|\n参考资料|\Z)",
        text,
        re.S,
    )
    if match:
        suggestion = _clean(match.group(1))[:5000]
    level_match = re.search(r"风险等级\s*([^\n]+)", text)
    return [
        {
            "漏洞名称": name,
            "风险等级": level_match.group(1) if level_match else "",
            "漏洞描述": _clean(text[:2000]),
            "加固建议": suggestion,
            "CVE编号": "、".join(cves),
        }
    ]


def _parse_legacy_doc(path: Path) -> list[dict]:
    """Convert binary .doc with an available local converter."""
    with tempfile.TemporaryDirectory() as directory:
        target_dir = Path(directory)
        libreoffice = shutil.which("soffice") or shutil.which("libreoffice")
        if libreoffice:
            subprocess.run(
                [
                    libreoffice,
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(target_dir),
                    str(path),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            converted = target_dir / f"{path.stem}.docx"
            return _parse_docx(converted)
        antiword = shutil.which("antiword")
        if antiword:
            result = subprocess.run(
                [antiword, str(path)],
                check=True,
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="ignore",
            )
            return _parse_labeled_text(result.stdout)
    raise RuntimeError(
        "Legacy DOC parsing requires LibreOffice (soffice) or antiword"
    )


def _safe_extract_zip(path: Path, target: Path) -> list[Path]:
    files = []
    with zipfile.ZipFile(path) as archive:
        for member in archive.infolist():
            member_path = (target / member.filename).resolve()
            if target.resolve() not in member_path.parents and member_path != target.resolve():
                continue
            if member.is_dir():
                continue
            member_path.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, member_path.open("wb") as destination:
                shutil.copyfileobj(source, destination)
            files.append(member_path)
    return files


def _safe_extract_rar(path: Path, target: Path) -> list[Path]:
    try:
        import rarfile
    except ImportError as exc:
        raise RuntimeError("RAR parsing requires rarfile and an unrar backend") from exc
    files = []
    with rarfile.RarFile(path) as archive:
        for member in archive.infolist():
            member_path = (target / member.filename).resolve()
            if target.resolve() not in member_path.parents and member_path != target.resolve():
                continue
            if member.isdir():
                continue
            member_path.parent.mkdir(parents=True, exist_ok=True)
            with archive.open(member) as source, member_path.open("wb") as destination:
                shutil.copyfileobj(source, destination)
            files.append(member_path)
    return files


def _parse_jianshi_html(text: str) -> list[dict]:
    """Parse 坚石诚信 (Jet Sreality) HTML scan report text.

    Pattern: vulnerability name followed by a numeric count, then
    "漏洞描述" / "解决办法" sections.
    """
    known_vulns = {
        "CORS", "SameSite", "Cookie", "HttpOnly", "Secure", "跨域",
        "XSS", "SQL注入", "CSRF", "信息泄露", "ClickJacking", "点击劫持",
        "弱口令", "明文传输", "不安全的HTTP方法", "目录遍历", "路径穿越",
        "文件包含", "命令执行", "未授权", "越权", "敏感信息泄露",
        "安全配置错误", "不安全设计", "注入", "过时的组件", "自带缺陷",
        "SRI", "子资源完整性", "域名访问限制", "用户认证信息",
        "密码表单自动完成", "电子邮箱", "应用错误信息", "HTML信息泄露",
    }
    lines = text.split("\n")
    rows = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        # Skip empty lines, numbers, and section headers
        if not line or re.match(r"^\d+$", line) or line in ("漏洞描述", "解决办法", "漏洞名称", "风险等级", "漏洞链接"):
            i += 1
            continue
        # Check if this looks like a vulnerability name (contains known keywords or is a reasonable length)
        has_kw = any(kw in line for kw in known_vulns)
        if not has_kw and (len(line) < 4 or len(line) > 60):
            i += 1
            continue
        # Skip common non-vuln lines
        if line in ("漏洞（种", "实例（个", "插件ID", "风险等级图标", "比较危险"):
            i += 1
            continue
        # This might be a vulnerability name. Look ahead for 漏洞描述/解决办法
        desc = ""
        fix = ""
        level = ""
        j = i + 1
        while j < min(i + 100, len(lines)):
            lj = lines[j].strip()
            if lj == "漏洞描述":
                k = j + 1
                desc_parts = []
                while k < min(j + 30, len(lines)):
                    lk = lines[k].strip()
                    if lk in ("解决办法", "漏洞描述", "漏洞名称", "风险等级", "漏洞链接") or (re.match(r"^\d+$", lk) and len(lk) < 5):
                        break
                    if lk:
                        desc_parts.append(lk)
                    k += 1
                desc = " ".join(desc_parts)
                j = k
                continue
            if lj == "解决办法":
                k = j + 1
                fix_parts = []
                while k < min(j + 30, len(lines)):
                    lk = lines[k].strip()
                    if lk in ("解决办法", "漏洞描述", "漏洞名称", "风险等级", "漏洞链接") or (re.match(r"^\d+$", lk) and len(lk) < 5):
                        break
                    if lk:
                        fix_parts.append(lk)
                    k += 1
                fix = " ".join(fix_parts)
                j = k
                continue
            j += 1
        if line:
            rows.append({"漏洞名称": line, "漏洞描述": desc, "加固建议": fix, "风险等级": "medium"})
            i = j
        else:
            i += 1
    return rows


def _parse_labeled_text(text: str) -> list[dict]:
    blocks = re.split(r"\n(?=(?:【?(?:严重|高危|中危|低危|信息)】?)?\s*\d*[.、]?\s*[^。\n]{2,60})", text)
    rows = []
    for block in blocks:
        name = re.search(r"(?:漏洞名称|问题名称|标题)[：:]\s*([^\n]+)", block)
        if not name:
            name = re.search(r"【(严重|高危|中危|低危|信息)】\s*([^\n]+)", block)
        if not name:
            continue
        title = name.group(2) if name.lastindex and name.lastindex >= 2 else name.group(1)
        row = {"漏洞名称": title}
        for field, labels in {
            "风险等级": ("风险等级", "漏洞等级"),
            "漏洞描述": ("漏洞描述", "问题描述"),
            "漏洞危害": ("漏洞危害", "影响"),
            "加固建议": ("加固建议", "修复建议", "整改建议", "解决方案"),
        }.items():
            pattern = "|".join(re.escape(label) for label in labels)
            match = re.search(
                rf"(?:{pattern})[：:]\s*(.*?)(?=\n(?:漏洞描述|问题描述|漏洞危害|影响|加固建议|修复建议|整改建议|解决方案)[：:]|\Z)",
                block,
                re.S,
            )
            row[field] = _clean(match.group(1)) if match else ""
        rows.append(row)
    return rows


def parse_report(path: str | Path) -> dict:
    report = Path(path)
    suffix = report.suffix.lower()
    if suffix in (".xlsx", ".xls"):
        rows = _parse_excel(report)
    elif suffix in (".csv", ".tsv"):
        rows = _parse_delimited(report)
    elif suffix == ".docx":
        rows = _parse_docx(report)
    elif suffix == ".doc":
        rows = _parse_legacy_doc(report)
    elif suffix == ".pdf":
        rows = _parse_pdf(report)
    elif suffix == ".json":
        data = json.loads(report.read_text(encoding="utf-8"))
        if isinstance(data, dict):
            rows = data.get("vulns") or data.get("items") or data.get("results") or [data]
        else:
            rows = data
    elif suffix in (".html", ".htm"):
        try:
            from bs4 import BeautifulSoup
        except ImportError as exc:
            raise RuntimeError("HTML parsing requires beautifulsoup4") from exc
        text = BeautifulSoup(report.read_text(encoding="utf-8", errors="ignore"), "html.parser").get_text("\n")
        rows = _parse_jianshi_html(text) or _parse_labeled_text(text)
    elif suffix in (".txt", ".md", ".log", ".out", ".properties"):
        rows = _parse_labeled_text(report.read_text(encoding="utf-8", errors="ignore"))
    elif suffix in (".zip", ".rar"):
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory)
            extracted = (
                _safe_extract_zip(report, target)
                if suffix == ".zip"
                else _safe_extract_rar(report, target)
            )
            vulns = []
            parse_errors = []
            for extracted_file in extracted:
                try:
                    parsed = parse_report(extracted_file)
                    for vuln in parsed["vulns"]:
                        vuln["source_file"] = (
                            f"{report.name}!{extracted_file.relative_to(target)}"
                        )
                    vulns.extend(parsed["vulns"])
                except Exception as exc:
                    parse_errors.append(
                        {
                            "file": str(extracted_file.relative_to(target)),
                            "error": str(exc),
                        }
                    )
            for index, vuln in enumerate(vulns, 1):
                vuln["id"] = index
            return {
                "source_file": str(report),
                "format": suffix.lstrip("."),
                "total": len(vulns),
                "vulns": vulns,
                "archive_parse_errors": parse_errors,
            }
    else:
        raise ValueError(f"Unsupported report format: {suffix or '(none)'}")

    vulns = normalize_rows(rows, report.name)
    for index, vuln in enumerate(vulns, 1):
        vuln["id"] = index
    return {
        "source_file": str(report),
        "format": suffix.lstrip("."),
        "total": len(vulns),
        "vulns": vulns,
    }
