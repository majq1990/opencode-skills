#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能文档生成器 - 支持 Markdown 和 DOCX 输出
默认输出 DOCX 格式
"""

import os
import sys
import json
import argparse
import re
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: python-docx 未安装，将只支持 Markdown 输出")
    print("安装命令: pip install python-docx")


class DocumentGenerator:
    """文档生成器类 - 支持 MD 和 DOCX"""
    
    def __init__(self, config_dir=None, templates_dir=None):
        """初始化文档生成器"""
        # 确定 config 目录
        if config_dir is None:
            script_dir = Path(__file__).parent
            self.config_dir = script_dir.parent / "config"
        else:
            self.config_dir = Path(config_dir)
        
        self.templates_dir = Path(templates_dir) if templates_dir else self.config_dir.parent / "templates"
        
        # 加载配置
        self.document_types = self.load_json("document-types.json")
        self.template_structure = self.load_json("template-structure.json")
        
    def load_json(self, filename):
        """加载JSON配置文件"""
        filepath = self.config_dir / filename
        if filepath.exists():
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def generate_document(self, doc_type, input_data, output_path=None, format_type="docx"):
        """
        生成文档
        
        Args:
            doc_type: 文档类型
            input_data: 输入数据字典
            output_path: 输出文件路径
            format_type: 输出格式 ("docx" 或 "md")
        """
        # 检查文档类型是否支持
        if doc_type not in self.document_types.get("documentTypes", {}):
            print(f"错误: 不支持的文档类型 '{doc_type}'")
            return False
        
        doc_info = self.document_types["documentTypes"][doc_type]
        
        # 生成文档内容
        if format_type == "docx":
            if not HAS_DOCX:
                print("错误: python-docx 未安装，无法生成 DOCX")
                print("安装命令: pip install python-docx")
                return False
            
            content = self.generate_content(doc_type, input_data)
            
            # 默认输出路径
            if output_path is None:
                output_path = f"{doc_info['name']}_{input_data.get('systemName', '系统')}_v{input_data.get('version', '1.0')}.docx"
            elif not output_path.endswith('.docx'):
                output_path = output_path.rsplit('.', 1)[0] + '.docx'
            
            self.save_as_docx(content, output_path, doc_info, input_data)
        else:
            content = self.generate_content_md(doc_type, input_data)
            
            if output_path is None:
                output_path = f"{doc_info['name']}_{input_data.get('systemName', '系统')}_v{input_data.get('version', '1.0')}.md"
            elif not output_path.endswith('.md'):
                output_path = output_path.rsplit('.', 1)[0] + '.md'
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
        
        print(f"文档已生成: {output_path}")
        return True
    
    def generate_content(self, doc_type, input_data):
        """生成文档内容结构（用于 DOCX）"""
        doc_info = self.document_types["documentTypes"][doc_type]
        
        return {
            "title": f"{input_data.get('systemName', '系统')} - {doc_info['name']}",
            "doc_id": doc_info['id'],
            "system_name": input_data.get('systemName', '系统'),
            "version": input_data.get('version', 'v1.0'),
            "author": input_data.get('author', '开发团队'),
            "date": datetime.now().strftime('%Y年%m月%d日'),
            "system_url": input_data.get('systemUrl', ''),
            "chapters": doc_info.get("chapters", []),
            "input_data": input_data,
            "doc_info": doc_info
        }
    
    def generate_content_md(self, doc_type, input_data):
        """生成 Markdown 内容"""
        doc_info = self.document_types["documentTypes"][doc_type]
        content = []
        
        # 文档标题
        title = f"{input_data.get('systemName', '系统')} - {doc_info['name']}"
        content.append(f"# {title}\n")
        content.append("---\n")
        
        # 文档信息
        content.append("## 文档信息\n")
        content.append("| 项目 | 内容 |\n")
        content.append("|------|------|\n")
        content.append(f"| **产品名称** | {input_data.get('systemName', '系统')} |\n")
        content.append(f"| **文档编号** | {doc_info['id']} |\n")
        content.append(f"| **文档版本** | {input_data.get('version', 'v1.0')} |\n")
        content.append(f"| **编写日期** | {datetime.now().strftime('%Y年%m月%d日')} |\n")
        if 'author' in input_data:
            content.append(f"| **编写人** | {input_data['author']} |\n")
        content.append("\n")
        
        # 修订历史
        content.append("## 修订历史\n")
        content.append("| 版本 | 修订日期 | 修订人 | 修订内容 |\n")
        content.append("|------|----------|--------|----------|\n")
        content.append(f"| {input_data.get('version', 'v1.0')} | {datetime.now().strftime('%Y-%m-%d')} | {input_data.get('author', '开发团队')} | 初始版本 |\n")
        content.append("\n")
        
        # 目录
        content.append("## 目录\n")
        for chapter in doc_info.get("chapters", []):
            content.append(f"- [{chapter}](#{chapter})\n")
        content.append("\n---\n\n")
        
        # 各章节
        for chapter in doc_info.get("chapters", []):
            content.extend(self.generate_chapter_md(chapter, input_data))
            content.append("\n")
        
        # 附录
        content.append("## 附录\n\n")
        content.append("---\n\n**文档结束**\n")
        
        return "".join(content)
    
    def generate_chapter_md(self, chapter, input_data):
        """生成 Markdown 章节"""
        content = []
        content.append(f"## {chapter}\n\n")
        
        if "引言" in chapter:
            content.append("### 编写目的\n\n")
            content.append(f"本文档旨在详细说明{input_data.get('systemName', '系统')}的技术细节。\n\n")
            content.append("### 背景\n\n")
            content.append(f"{input_data.get('systemName', '系统')}是一个{input_data.get('description', '企业级应用系统')}。\n\n")
        elif "概述" in chapter:
            content.append("### 系统目标\n\n")
            content.append("- 提供高效、稳定的系统服务\n")
            content.append("- 支持多用户并发访问\n\n")
        elif "架构" in chapter:
            content.append("### 总体架构\n\n")
            content.append("系统采用分层架构设计。\n\n")
        else:
            content.append("（待补充内容）\n\n")
        
        return content
    
    def save_as_docx(self, content, output_path, doc_info, input_data):
        """保存为 DOCX 格式"""
        doc = Document()
        
        # 设置中文字体
        self._set_chinese_font(doc)
        
        # 封面标题
        title = doc.add_paragraph()
        title_run = title.add_run(content["title"])
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 文档信息表格
        doc.add_heading('文档信息', level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("产品名称", content["system_name"]),
            ("文档编号", content["doc_id"]),
            ("文档版本", content["version"]),
            ("编写日期", content["date"]),
            ("编写人", content["author"])
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
        
        doc.add_paragraph()
        
        # 修订历史
        doc.add_heading('修订历史', level=1)
        history_table = doc.add_table(rows=2, cols=4)
        history_table.style = 'Table Grid'
        
        headers = ["版本", "修订日期", "修订人", "修订内容"]
        for i, h in enumerate(headers):
            history_table.cell(0, i).text = h
        
        history_table.cell(1, 0).text = content["version"]
        history_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d')
        history_table.cell(1, 2).text = content["author"]
        history_table.cell(1, 3).text = "初始版本"
        
        doc.add_paragraph()
        
        # 目录
        doc.add_heading('目录', level=1)
        for chapter in content["chapters"]:
            p = doc.add_paragraph(chapter, style='List Bullet')
        
        doc.add_page_break()
        
        # 各章节内容
        for chapter in content["chapters"]:
            self._add_chapter(doc, chapter, input_data)
        
        # 附录
        doc.add_heading('附录', level=1)
        doc.add_paragraph("（根据实际情况补充附录内容）")
        
        # 保存
        doc.save(output_path)
    
    def _set_chinese_font(self, doc):
        """设置中文字体"""
        pass  # python-docx 会使用系统默认字体
    
    def _add_chapter(self, doc, chapter, input_data):
        """添加章节"""
        doc.add_heading(chapter, level=1)
        
        # 根据章节类型添加内容
        if "引言" in chapter:
            doc.add_heading('编写目的', level=2)
            doc.add_paragraph(f"本文档旨在详细说明{input_data.get('systemName', '系统')}的技术架构、功能模块等关键技术细节，为系统开发、测试、部署和维护提供技术指导。")
            
            doc.add_heading('背景', level=2)
            doc.add_paragraph(f"{input_data.get('systemName', '系统')}是一个{input_data.get('description', '企业级应用系统')}，采用现代化的技术架构，为用户提供高效、可靠的服务。")
            
            doc.add_heading('术语定义', level=2)
            term_table = doc.add_table(rows=4, cols=2)
            term_table.style = 'Table Grid'
            terms = [
                ("术语", "定义"),
                ("API", "Application Programming Interface, 应用程序接口"),
                ("REST", "Representational State Transfer, 表述性状态转移"),
                ("BPMN", "Business Process Model and Notation, 业务流程建模与表示法")
            ]
            for i, (term, defn) in enumerate(terms):
                term_table.cell(i, 0).text = term
                term_table.cell(i, 1).text = defn
        
        elif "概述" in chapter and "系统" in chapter:
            doc.add_heading('系统目标', level=2)
            doc.add_paragraph("• 提供高效、稳定的系统服务")
            doc.add_paragraph("• 支持多用户并发访问")
            doc.add_paragraph("• 确保数据安全和系统可靠性")
            
            doc.add_heading('系统特点', level=2)
            features = input_data.get('features', [])
            if features:
                for f in features:
                    doc.add_paragraph(f"• {f}")
            else:
                doc.add_paragraph("（待补充系统特点）")
        
        elif "架构" in chapter:
            doc.add_heading('总体架构', level=2)
            doc.add_paragraph("系统采用分层架构设计，主要包括表现层、业务层和数据层。")
            
            doc.add_heading('技术栈', level=2)
            tech = input_data.get('techStack', {})
            if tech:
                tech_table = doc.add_table(rows=len(tech)+1, cols=2)
                tech_table.style = 'Table Grid'
                tech_table.cell(0, 0).text = "层次"
                tech_table.cell(0, 1).text = "技术"
                for i, (k, v) in enumerate(tech.items(), 1):
                    tech_table.cell(i, 0).text = k
                    tech_table.cell(i, 1).text = str(v)
        
        elif "功能" in chapter:
            doc.add_heading('功能模块列表', level=2)
            features = input_data.get('features', [])
            if features:
                for i, f in enumerate(features, 1):
                    doc.add_heading(f'{i}. {f}', level=3)
                    doc.add_paragraph("（待补充功能详细说明）")
            else:
                doc.add_paragraph("（待补充功能模块）")
        
        elif "数据库" in chapter:
            doc.add_heading('数据库概述', level=2)
            db_type = input_data.get('databaseType', '关系型数据库')
            doc.add_paragraph(f"系统使用 {db_type} 作为数据库。")
            
            doc.add_heading('表结构说明', level=2)
            tables = input_data.get('tables', [])
            if tables:
                table_doc = doc.add_table(rows=len(tables)+1, cols=2)
                table_doc.style = 'Table Grid'
                table_doc.cell(0, 0).text = "表名"
                table_doc.cell(0, 1).text = "说明"
                for i, t in enumerate(tables, 1):
                    table_doc.cell(i, 0).text = t.get('name', '')
                    table_doc.cell(i, 1).text = t.get('description', '')
            else:
                doc.add_paragraph("（待补充表结构信息）")
        
        elif "接口" in chapter:
            doc.add_heading('接口概述', level=2)
            doc.add_paragraph("系统提供REST API接口，支持HTTP/JSON协议。")
            
            base_url = input_data.get('baseUrl', '')
            if base_url:
                doc.add_paragraph(f"Base URL: {base_url}")
            
            doc.add_heading('主要接口', level=2)
            apis = input_data.get('apiDocs', [])
            if apis:
                for api in apis:
                    doc.add_heading(api.get('name', ''), level=3)
                    doc.add_paragraph(f"方法: {api.get('method', 'GET')}")
                    doc.add_paragraph(f"路径: {api.get('path', '')}")
                    doc.add_paragraph(f"说明: {api.get('description', '')}")
            else:
                doc.add_paragraph("（待补充接口信息）")
        
        elif "安全" in chapter:
            doc.add_heading('认证机制', level=2)
            doc.add_paragraph("• 用户名/密码认证")
            doc.add_paragraph("• Token认证")
            
            doc.add_heading('安全配置', level=2)
            doc.add_paragraph("• 使用HTTPS协议传输")
            doc.add_paragraph("• 密码加密存储")
            doc.add_paragraph("• SQL注入防护")
            doc.add_paragraph("• XSS防护")
        
        elif "性能" in chapter:
            doc.add_heading('性能指标', level=2)
            perf_table = doc.add_table(rows=4, cols=2)
            perf_table.style = 'Table Grid'
            perf_data = [
                ("指标", "目标值"),
                ("接口响应时间", "< 200ms"),
                ("并发用户数", "> 1000"),
                ("系统可用性", "> 99.9%")
            ]
            for i, (k, v) in enumerate(perf_data):
                perf_table.cell(i, 0).text = k
                perf_table.cell(i, 1).text = v
        
        elif "部署" in chapter:
            doc.add_heading('部署架构', level=2)
            doc.add_paragraph("系统支持多种部署方式：")
            doc.add_paragraph("1. 单机部署")
            doc.add_paragraph("2. 集群部署")
            doc.add_paragraph("3. Docker容器部署")
        
        elif "环境" in chapter:
            doc.add_heading('硬件要求', level=2)
            hw_table = doc.add_table(rows=4, cols=2)
            hw_table.style = 'Table Grid'
            hw_data = [
                ("配置项", "要求"),
                ("CPU", "4核+"),
                ("内存", "8GB+"),
                ("磁盘", "100GB+")
            ]
            for i, (k, v) in enumerate(hw_data):
                hw_table.cell(i, 0).text = k
                hw_table.cell(i, 1).text = v
            
            doc.add_heading('系统访问', level=2)
            if input_data.get('systemUrl'):
                doc.add_paragraph(f"系统地址: {input_data['systemUrl']}")
        
        else:
            doc.add_paragraph("（待补充内容）")
        
        doc.add_paragraph()


def convert_md_to_docx(md_path, docx_path=None):
    """将 Markdown 文件转换为 DOCX"""
    if not HAS_DOCX:
        print("错误: python-docx 未安装")
        return False
    
    if docx_path is None:
        docx_path = md_path.rsplit('.', 1)[0] + '.docx'
    
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    doc = Document()
    
    lines = md_content.split('\n')
    i = 0
    in_table = False
    table_data = []
    
    while i < len(lines):
        line = lines[i]
        
        # 标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # 表格
        elif line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_data = []
            
            # 解析表格行
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if not all(c.replace('-', '').replace(':', '') == '' for c in cells):  # 非分隔行
                table_data.append(cells)
        else:
            if in_table and table_data:
                # 创建表格
                if len(table_data) > 0:
                    num_cols = len(table_data[0])
                    num_rows = len(table_data)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for r_idx, row_data in enumerate(table_data):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < num_cols:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = cell_text
                                # 加粗表头
                                if r_idx == 0:
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.font.bold = True
                
                in_table = False
                table_data = []
            
            # 普通段落
            if line.strip():
                # 处理列表
                if line.strip().startswith('- ') or line.strip().startswith('* '):
                    doc.add_paragraph(line.strip()[2:], style='List Bullet')
                elif re.match(r'^\d+\.\s', line.strip()):
                    doc.add_paragraph(re.sub(r'^\d+\.\s', '', line.strip()), style='List Number')
                else:
                    p = doc.add_paragraph()
                    # 处理粗体和斜体
                    text = line.strip()
                    # 简单处理粗体
                    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                    text = re.sub(r'\*(.+?)\*', r'\1', text)
                    p.add_run(text)
            elif line.strip() == '':
                pass  # 空行
        
        i += 1
    
    # 处理最后的表格
    if in_table and table_data:
        num_cols = len(table_data[0])
        num_rows = len(table_data)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        for r_idx, row_data in enumerate(table_data):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < num_cols:
                    table.cell(r_idx, c_idx).text = cell_text
    
    doc.save(docx_path)
    print(f"已转换: {md_path} -> {docx_path}")
    return True


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='智能文档生成器 (支持 DOCX 输出)')
    parser.add_argument('--type', help='文档类型')
    parser.add_argument('--input', help='输入数据文件(JSON格式)')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--format', default='docx', choices=['docx', 'md'], help='输出格式 (默认: docx)')
    parser.add_argument('--list-types', action='store_true', help='列出支持的文档类型')
    parser.add_argument('--convert', help='将 Markdown 文件转换为 DOCX')
    
    args = parser.parse_args()
    
    generator = DocumentGenerator()
    
    # 列出支持的文档类型
    if args.list_types:
        print("支持的文档类型:")
        print("-" * 50)
        for doc_type, info in generator.document_types.get("documentTypes", {}).items():
            print(f"{info['id']}: {doc_type}")
            print(f"  描述: {info['description']}")
            print()
        return 0
    
    # 转换 MD 到 DOCX
    if args.convert:
        convert_md_to_docx(args.convert, args.output)
        return 0
    
    # 检查必需参数
    if not args.type or not args.input:
        parser.print_help()
        return 1
    
    # 读取输入数据
    with open(args.input, 'r', encoding='utf-8') as f:
        input_data = json.load(f)
    
    # 生成文档
    success = generator.generate_document(args.type, input_data, args.output, args.format)
    
    return 0 if success else 1


if __name__ == "__main__":
    sys.exit(main())